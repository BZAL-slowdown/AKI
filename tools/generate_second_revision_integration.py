from __future__ import annotations

import shutil
from pathlib import Path

import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(r"D:\AKI")
OUT = ROOT / "second_revision_integrated"


def set_para_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def insert_paragraph_after(paragraph, text: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    return new_para


def find_para(doc: Document, starts: str | None = None, contains: str | None = None, exact: str | None = None):
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if exact is not None and text == exact:
            return i, para
        if starts is not None and text.startswith(starts):
            return i, para
        if contains is not None and contains in text:
            return i, para
    raise ValueError(f"Paragraph not found: starts={starts!r} contains={contains!r} exact={exact!r}")


def add_para(doc: Document, text: str = "", bold_label: str | None = None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    if bold_label:
        run = para.add_run(bold_label)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    if text:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    return para


def doc_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def build_response_sections():
    sections = []

    def add_section(title, comment, response, actions, location):
        sections.append(
            {
                "title": title,
                "comment": comment,
                "response": response,
                "actions": actions,
                "location": location,
            }
        )

    add_section(
        "Major Comment 1. External validity and overfitting",
        "The study is based on a single-center retrospective cohort with a limited sample size (n=194). Despite the use of few-shot learning, the complexity of the proposed architecture raises concerns regarding overfitting. The absence of external or temporal validation significantly limits generalizability.",
        "We thank the reviewer for emphasizing this central limitation. We agree that the single-center retrospective design and limited sample size restrict external validity. We therefore revised the manuscript to frame the work as an exploratory, internally validated model-development study rather than a model intended for direct clinical use. We now explicitly state that no external or temporal validation cohort was available for this revision and that multicenter, temporal, and prospective validation will be required before the model can be generalized or used in routine clinical decision-making.\n\nTo address overfitting concerns more transparently, we added a reproducible second-revision analysis script and additional internal validation evidence. The newly added tabular baseline analysis began with 194 records, excluded 1 row because sequence number or postoperative AKI label was missing, and analyzed 193 patients (132 AKI-positive and 61 AKI-negative). Classical baseline models were evaluated using five-fold stratified outer cross-validation with three-fold inner hyperparameter selection. The best classical clinical baseline was gradient boosting, with AUC 0.636 (bootstrap 95% CI 0.548-0.718), Brier score 0.223, and MCC 0.172. In repeated stratified five-fold cross-validation with five repeats, gradient boosting showed mean AUC 0.628 +/- 0.084 SD, while random forest showed mean AUC 0.620 +/- 0.087 SD. These findings are reported as internal evidence only and are used to temper interpretation rather than to claim external generalizability.",
        "We revised the Abstract, Methods, Results, and Discussion to state that the model evidence is internal cross-validation evidence only. We also added classical baseline and repeated-CV robustness results to the supplementary materials.",
        "Abstract; Methodology - Model Evaluation; Results - Model Performance; Discussion - limitations and clinical interpretation; S1-S3 Tables.",
    )
    add_section(
        "Major Comment 2. Incomplete clinical variables",
        "The model relies exclusively on preoperative variables. Key intraoperative determinants of postoperative AKI (warm ischemia time, pneumoperitoneum pressure, intraoperative hypotension, blood loss) are not included. Some selected variables have weak physiological links to AKI.",
        "We agree with the reviewer. Warm ischemia time, pneumoperitoneum pressure and duration, intraoperative hypotension burden, blood loss, operative duration, fluid balance, vasopressor exposure, and nephrotoxic exposure are clinically important determinants of postoperative AKI. These variables were not available in complete structured form in the retrospective dataset and were therefore not included in the current model. We revised the manuscript to make clear that the model is a preoperative risk-stratification model rather than a full perioperative dynamic prediction model. We also state that the absence of intraoperative variables may limit discrimination, calibration, threshold transportability, and clinical usefulness.",
        "We added explicit statements in the clinical-feature section and Discussion noting that intraoperative variables were unavailable/incomplete and were not modeled. We also softened clinical-utility language accordingly.",
        "Methodology - Clinical Feature Selection and Definition; Abstract; Discussion - limitations and future work.",
    )
    add_section(
        "Major Comment 3. Model complexity",
        "The accumulation of multiple advanced mechanisms (multimodal fusion, multi-prototype clustering, pseudo-sample generation, multiple loss functions) appears disproportionate to the sample size and may limit reproducibility and clinical translation.",
        "We appreciate this comment and agree that model complexity must be interpreted cautiously in a 194-patient retrospective cohort. We revised the manuscript so that the architecture is presented as an exploratory proof-of-concept design rather than definitive evidence of clinical superiority. The newly added classical baseline analyses provide context for this complexity: optimized clinical-only baselines had moderate internal discrimination, with gradient boosting achieving AUC 0.636 (95% CI 0.548-0.718) and repeated-CV AUC 0.628 +/- 0.084. These findings support a balanced interpretation: the proposed multimodal model may improve internal discrimination within this cohort, but its added complexity requires external validation, calibration assessment, and prospective testing before clinical translation.",
        "We condensed and tempered the interpretation of the model architecture, added simpler clinical baselines as benchmarks, and added explicit discussion that residual overfitting and reproducibility concerns remain.",
        "Methodology - model design/evaluation; Results - baseline comparison; Discussion - limitations.",
    )
    add_section(
        "Major Comment 4. Baseline model comparison",
        "The clinical-only model demonstrates poor performance. Comparisons with optimized classical models (e.g., penalized regression or gradient boosting) are lacking and would strengthen the manuscript.",
        "We agree and added optimized classical clinical baseline comparisons using the 193-patient tabular analytic set and 11 preoperative predictors. The analysis was implemented in second_revision_analysis.py. Missing clinical values were handled by median imputation within each training fold to avoid information leakage. Four classical baselines were evaluated: penalized logistic regression, random forest, gradient boosting, and support vector machine. Each model used five-fold stratified outer cross-validation with three-fold inner hyperparameter selection by ROC AUC.\n\nThe best classical baseline was gradient boosting, with AUC 0.636 (bootstrap 95% CI 0.548-0.718), average precision 0.795, Brier score 0.223, sensitivity 0.864, specificity 0.279, and MCC 0.172 at threshold 0.5. Random forest had AUC 0.617 (95% CI 0.526-0.700), Brier score 0.220, and MCC 0.152. Penalized logistic regression and support vector machine performed less well by AUC (0.480 and 0.565, respectively). In repeated stratified five-fold CV with five repeats, gradient boosting remained the best model by mean AUC (0.628 +/- 0.084), followed closely by random forest (0.620 +/- 0.087).",
        "We added a classical-baseline comparison paragraph to the Methods and Results and provide detailed performance metrics in S2 Table and repeated-CV robustness results in S3 Table.",
        "Methodology - Model Evaluation; Results - Model Performance; S2 Table; S3 Table.",
    )
    add_section(
        "Major Comment 5. Clinical evaluation: calibration, decision-curve analysis, and threshold",
        "Calibration metrics and decision-curve analysis are not provided. The clinical implications of the selected decision threshold are insufficiently discussed.",
        "We thank the reviewer for this important point. We added calibration and decision-curve analyses to the second-revision evidence package for the optimized classical clinical baselines. The performance table now reports Brier score, calibration intercept, calibration slope, and expected calibration error. Calibration curves are provided in S1 Fig, and decision-curve analysis is provided in S2 Fig.\n\nThe clinical interpretation has been tempered. For the gradient boosting baseline, Brier score was 0.223, calibration intercept was 0.413, calibration slope was 0.349, and expected calibration error was 0.109. Decision-curve analysis was included for transparency, but it should not be interpreted as establishing clinical utility. Across threshold probabilities from 0.10 to 0.50, gradient boosting did not clearly exceed the treat-all net-benefit reference. We therefore revised the manuscript to state that threshold selection remains preliminary and must be re-evaluated in external and prospective cohorts before any clinical workflow use.",
        "We added calibration metrics, calibration curves, decision-curve analysis, and a cautious threshold interpretation. We explicitly state that these are internal classical-baseline analyses, not final multimodal-model external validation and not evidence for deployment readiness.",
        "Methodology - Model Evaluation; Results - Model Performance; Discussion - clinical interpretation; S1 Fig; S2 Fig.",
    )
    add_section(
        "Minor Comments",
        "The manuscript remains lengthy and could benefit from further condensation. Claims regarding clinical utility should be further tempered. SHAP analysis is descriptive and should be cautiously interpreted. Figures and legends could be simplified.",
        "We agree with these points and revised the manuscript accordingly. We shortened and tempered several statements in the Abstract, Results, and Discussion; reframed SHAP as descriptive model-attribution analysis rather than causal or interventional evidence; and added supporting-information captions for the new tables and figures. We also prepared a supplementary materials package with clearly labeled S1-S3 Tables and S1-S4 Figures. The calibration and decision-curve figures are explicitly labeled as internal classical-baseline evidence only.",
        "Clinical-utility claims were softened, SHAP interpretation was made cautious, and supplementary figure/table captions were added. The submitted figure package should be checked for final numbering consistency before upload.",
        "Abstract; Results; SHAP-based Feature Importance; Discussion; Supporting information captions; supplementary materials package.",
    )
    return sections


def revise_manuscript(manuscript_out: Path) -> None:
    doc = Document(str(manuscript_out))

    abstract = (
        "To optimize preoperative decision-making, early prediction of acute kidney injury (AKI), a major postoperative complication, remains clinically important. "
        "This single-center retrospective study developed an exploratory multimodal few-shot prototype network using real-world data from 194 patients undergoing laparoscopic renal surgery. "
        "The model integrated 11 preoperative clinical parameters encoded by a multilayer perceptron with imaging features extracted by ResNet-50, using gated fusion and class-weighted learning to address small-sample and class-imbalance challenges. "
        "In internal five-fold cross-validation, the fused model achieved an area under the receiver operating characteristic curve (AUC) of 0.816, outperforming the clinical-only (AUC = 0.578) and imaging-only (AUC = 0.510) models within this cohort. "
        "Additional second-revision analyses compared optimized classical clinical baselines; gradient boosting showed the highest clinical-baseline AUC of 0.636 (bootstrap 95% CI 0.548-0.718) in internal cross-validation, with repeated stratified cross-validation yielding a mean AUC of 0.628 +/- 0.084. "
        "SHAP analysis provided a descriptive summary of feature contributions, highlighting preoperative creatinine, lactate, and age. "
        "Because the study lacks external or temporal validation and did not include intraoperative variables such as warm ischemia time, pneumoperitoneum pressure, intraoperative hypotension, or blood loss, the findings should be interpreted as hypothesis-generating internal-validation evidence requiring multicenter prospective validation before clinical deployment."
    )
    _, para = find_para(doc, starts="To optimize preoperative decision-making")
    set_para_text(para, abstract)

    contains_replacements = [
        (
            "significantly improving",
            "Our few-shot prototype network integrates preoperative static clinical variables and imaging features, showing improved internal cross-validated discrimination in a limited-data setting while still requiring external validation.",
        ),
        (
            "addressing",
            "Model development used real-world, single-center data from Sichuan Provincial Second People’s Hospital; therefore, the current findings primarily reflect the local cohort and require temporal and multicenter validation before broader generalization.",
        ),
        (
            "visualizing risk factors",
            "SHAP interpretability analysis provides a descriptive visualization of model feature contributions, which may help clinicians understand model behavior but should not be interpreted as causal evidence or evidence for clinical deployment.",
        ),
        (
            "geographically flexible postoperative AKI prediction",
            "In conclusion, this study provides an exploratory technical basis for multimodal preoperative AKI risk prediction in a resource-limited, single-center cohort. Prospective multicenter validation, inclusion of intraoperative variables, and local recalibration are needed before clinical use.",
        ),
    ]
    for needle, replacement in contains_replacements:
        for p in doc.paragraphs:
            if needle in p.text:
                set_para_text(p, replacement)
                break

    _, p = find_para(doc, starts="This study might more accurately determine")
    insert_paragraph_after(
        p,
        "For transparency in the second-revision analyses, the raw clinical spreadsheet contained 194 records. The newly added tabular classical-baseline analysis excluded one record because either the sequence number or postoperative AKI outcome label was missing, leaving 193 patients (132 postoperative AKI-positive and 61 postoperative AKI-negative) for the clinical-baseline comparison. Image availability was audited separately and is summarized in S1 Table.",
    )

    _, p = find_para(doc, starts="Furthermore, blood urea nitrogen")
    insert_paragraph_after(
        p,
        "The current model was designed as a preoperative risk-stratification model. Important intraoperative determinants of postoperative AKI, including warm ischemia time, pneumoperitoneum pressure and duration, intraoperative hypotension burden, estimated blood loss, operative duration, fluid balance, vasopressor exposure, and nephrotoxic exposure, were unavailable or incomplete in structured form in this retrospective dataset and were therefore not modeled. Their absence is acknowledged as a major limitation and a priority for future prospective data collection.",
    )

    _, p = find_para(doc, starts="To preserve sample size and minimize")
    insert_paragraph_after(
        p,
        "For the added classical-baseline analyses requested during peer review, missing clinical values were imputed within each training fold and standardization/model fitting were performed inside scikit-learn pipelines to avoid information leakage from validation folds.",
    )

    _, p = find_para(doc, starts="To evaluate our proposed multimodal model")
    set_para_text(
        p,
        "To evaluate our proposed multimodal model, we employed an internal nested cross-validation strategy with five-fold stratified partitioning to maintain balanced class distributions across folds. Fig 6 illustrates the evaluation workflow. This internal resampling approach was used to reduce bias from data partitioning and to assess model stability within the available single-center cohort; it should not be interpreted as external or temporal validation.",
    )
    insert_paragraph_after(
        p,
        "In response to peer-review requests, we added optimized classical clinical baseline comparisons using the 193-patient tabular analytic set and the same 11 preoperative clinical predictors. Penalized logistic regression, random forest, gradient boosting, and support vector machine models were evaluated using five-fold stratified outer cross-validation with three-fold inner hyperparameter selection by ROC AUC. Performance metrics included AUC, average precision, Brier score, calibration intercept and slope, expected calibration error, accuracy, balanced accuracy, precision, recall, F1 score, Matthews correlation coefficient (MCC), and confusion-matrix counts at threshold 0.5. Repeated stratified five-fold cross-validation with five repeats was used as an additional internal robustness check. Calibration curves and decision-curve analysis were generated for the optimized classical clinical baselines and are presented as internal baseline evidence only, not as external validation and not as evidence that the final multimodal model is clinically useful.",
    )

    _, p = find_para(doc, starts="Through systematic multimodal data integration")
    set_para_text(
        p,
        "Through systematic multimodal data integration, the fusion-based postoperative AKI prediction model showed higher internal cross-validated discrimination than the single-modality approaches within this cohort. The imaging-only model (ResNet-50) achieved an AUC of 0.510 in internal validation, the clinical-only model achieved an AUC of 0.578, and the fused multimodal model achieved an AUC of 0.816. These estimates should be interpreted as internal cross-validation results from a single-center retrospective cohort.",
    )
    _, p = find_para(doc, starts="Through SHAP-based interpretability analysis")
    set_para_text(
        p,
        "Through SHAP-based interpretability analysis, this study summarized the relative contributions of preoperative clinical variables to model outputs. These descriptive explanations may improve transparency, but they should not be interpreted as causal mechanisms or as sufficient evidence for clinical implementation.",
    )
    _, p = find_para(doc, starts="Although the representativeness")
    set_para_text(
        p,
        "Although the representativeness of the sample, the transportability of thresholds, and the external validity of the model require further improvement, the multimodal framework provides exploratory evidence for preoperative AKI risk stratification. Future work should prioritize multicenter and temporal validation, prospective collection of intraoperative variables, calibration assessment, and decision-curve evaluation before clinical deployment.",
    )
    _, p = find_para(doc, starts="Using five-fold cross-validation")
    set_para_text(
        p,
        "Using five-fold internal cross-validation, the postoperative acute kidney injury (AKI) prediction model developed in this study achieved an average AUC of 0.816 +/- 0.036. The overall ROC curve (Fig 7) remained above the diagonal reference line, indicating promising internal discrimination within this cohort. However, because the study lacks external or temporal validation, this result should be considered exploratory rather than evidence of generalizable clinical performance.",
    )
    _, p = find_para(doc, starts="The optimal threshold determined by Youden")
    set_para_text(
        p,
        "The optimal threshold determined by Youden’s index was 0.332. This internally derived threshold balanced sensitivity and specificity in the available cohort, but it is not necessarily transportable to other institutions, surgical workflows, or AKI prevalences. Future studies should re-estimate and recalibrate thresholds in external and prospective cohorts before clinical use.",
    )
    baseline = insert_paragraph_after(
        p,
        "Classical clinical baseline comparisons provided additional context (S2 Table). Among the optimized preoperative clinical baselines, gradient boosting achieved the highest internal cross-validated AUC of 0.636 (bootstrap 95% CI 0.548-0.718), average precision of 0.795, Brier score of 0.223, and MCC of 0.172 at threshold 0.5. Random forest achieved AUC 0.617 (95% CI 0.526-0.700), Brier score 0.220, and MCC 0.152, whereas penalized logistic regression and support vector machine achieved AUCs of 0.480 and 0.565, respectively. Repeated stratified five-fold cross-validation with five repeats showed similar results, with gradient boosting yielding mean AUC 0.628 +/- 0.084 and random forest yielding mean AUC 0.620 +/- 0.087 (S3 Table). These internal baseline results provide a benchmark but do not replace external validation.",
    )
    insert_paragraph_after(
        baseline,
        "Calibration and decision-curve analyses were added for the optimized classical clinical baselines (S1 Fig and S2 Fig). For the best classical baseline, gradient boosting, the Brier score was 0.223, calibration intercept was 0.413, calibration slope was 0.349, and expected calibration error was 0.109. Decision-curve analysis was included to transparently assess potential clinical net benefit; however, across threshold probabilities from 0.10 to 0.50, gradient boosting did not clearly exceed the treat-all reference. These findings support cautious interpretation and emphasize that clinical usefulness remains uncertain without external validation and prospective workflow evaluation.",
    )

    _, p = find_para(doc, starts="Preoperative serum creatinine")
    set_para_text(
        p,
        "Preoperative serum creatinine: the most prominent contributor in the SHAP summary, consistent with the importance of baseline renal function in postoperative AKI risk. This should be interpreted as model-attribution evidence rather than causal evidence.",
    )
    _, p = find_para(doc, starts="Overall, SHAP analysis")
    set_para_text(
        p,
        "Overall, SHAP analysis provided a descriptive summary of model feature contributions and highlighted clinically plausible variables such as creatinine, lactate, and age. These results may help readers understand model behavior, but they do not establish causality or define clinical intervention targets by themselves.",
    )
    _, p = find_para(doc, starts="This study's postoperative AKI prediction")
    set_para_text(
        p,
        "This study explores a multimodal few-shot approach for postoperative AKI prediction in a small, imbalanced, single-center retrospective cohort. The proposed architecture is methodologically exploratory and should be interpreted with caution because the combination of multimodal fusion, multi-prototype representation, pseudo-sample generation, and multiple losses may increase overfitting risk and reduce reproducibility when sample size is limited.",
    )
    _, p = find_para(doc, starts="Nonetheless, there are still certain restrictions")
    set_para_text(
        p,
        "Several important limitations must be emphasized. First, the study was retrospective and single-center, and no external or temporal validation cohort was available; therefore, all reported performance estimates are internal cross-validation results. Second, although few-shot learning, regularization, augmentation, and nested cross-validation may reduce overfitting risk, they cannot eliminate residual overfitting in a 194-patient cohort. Third, the current model provides static preoperative prediction and excludes key intraoperative and postoperative dynamic variables such as warm ischemia time, pneumoperitoneum pressure and duration, intraoperative hypotension, blood loss, operative duration, fluid balance, vasopressor exposure, and nephrotoxic exposure.",
    )
    _, p = find_para(doc, starts="Additionally, the temporal validity")
    set_para_text(
        p,
        "Additionally, temporal validity should be acknowledged. Because the model was developed from historical single-center data, changes in perioperative workflows, surgical techniques, anesthesia and hemodynamic management, laboratory testing, and AKI prevention protocols may alter model calibration and threshold performance over time. Future studies should include temporal validation, external multicenter cohorts, and local recalibration before clinical use.",
    )
    _, p = find_para(doc, starts="From a clinical perspective")
    set_para_text(
        p,
        "From a clinical perspective, the potential usefulness of the model lies in supporting future preoperative risk-stratification research rather than immediate clinical deployment. The added calibration and decision-curve analyses for optimized classical clinical baselines showed that clinical net benefit remains uncertain, reinforcing the need for prospective validation, decision-curve assessment of the final multimodal model, and workflow studies before any recommendation for routine clinical implementation.",
    )

    _, fig10 = find_para(doc, exact="Fig 10. Feature Importance for Postoperative AKI Prediction")
    cur = fig10
    for text in [
        "S1 Table. Case-flow and data-availability audit for the second-revision internal-validation analyses. The table summarizes raw rows, excluded rows, the final tabular baseline analytic set, outcome counts, image-file availability, YOLO annotation availability, clinical records without matching JPG files, and JPG files without TXT annotations.",
        "S2 Table. Internal cross-validated performance of optimized classical clinical baseline models. Penalized logistic regression, random forest, gradient boosting, and support vector machine were evaluated using the 193-patient tabular analytic set and 11 preoperative clinical predictors. Metrics include AUC with bootstrap 95% CI, average precision, Brier score, calibration metrics, threshold-based classification metrics, and confusion-matrix counts at threshold 0.5.",
        "S3 Table. Repeated stratified 5-fold cross-validation robustness summary for optimized classical clinical baselines. The table reports mean, standard deviation, and approximate 95% confidence intervals across 25 validation folds for discrimination, classification, MCC, and Brier-score metrics.",
        "S1 Fig. Internal cross-validated calibration curves for optimized classical clinical baseline models. Curves compare predicted and observed postoperative AKI risk for classical clinical baselines. These curves provide internal calibration evidence for baseline models only and should not be interpreted as final multimodal-model calibration.",
        "S2 Fig. Internal decision-curve analysis for optimized classical clinical baseline models. Net benefit is plotted across threshold probabilities for classical clinical baselines with treat-all and treat-none references. This internal analysis is included for transparency and does not establish clinical utility or readiness for deployment.",
        "S3 Fig. Receiver operating characteristic curves for optimized classical clinical baseline models under internal cross-validation. This optional figure visualizes discrimination of the classical clinical baselines and complements S2 Table.",
        "S4 Fig. Precision-recall curves for optimized classical clinical baseline models under internal cross-validation. This optional figure is useful because the analytic set is outcome-imbalanced.",
    ]:
        cur = insert_paragraph_after(cur, text)
    doc.save(str(manuscript_out))


def revise_supporting_info(support_out: Path) -> None:
    doc = Document(str(support_out))
    doc.add_page_break()
    add_para(
        doc,
        "S8 Appendix. Second-revision baseline comparison, calibration, decision-curve analysis, and robustness checks",
        bold_label="",
    ).runs[0].bold = True
    for text in [
        "In response to reviewer requests, additional optimized classical clinical baseline analyses were performed using the same preoperative clinical-feature set. The raw clinical spreadsheet contained 194 records; one record was excluded because either the sequence number or postoperative AKI outcome label was missing, leaving 193 patients for the tabular baseline analyses (132 postoperative AKI-positive and 61 postoperative AKI-negative cases).",
        "The evaluated classical clinical baselines were penalized logistic regression, random forest, gradient boosting, and support vector machine models. Missing clinical values were imputed within each training fold. Models were evaluated using five-fold stratified outer cross-validation with three-fold inner hyperparameter selection by ROC AUC. Repeated stratified five-fold cross-validation with five repeats was used as an additional internal robustness check.",
        "The best classical clinical baseline was gradient boosting, with AUC 0.636 (bootstrap 95% CI 0.548-0.718), Brier score 0.223, and MCC 0.172 at threshold 0.5. Repeated cross-validation showed a mean AUC of 0.628 +/- 0.084 for gradient boosting and 0.620 +/- 0.087 for random forest. These analyses are internal validation evidence only and do not replace external or temporal validation.",
        "Calibration curves and decision-curve analyses were generated for the optimized classical clinical baselines. The decision-curve results should be interpreted cautiously and should not be described as establishing clinical utility or deployment readiness.",
    ]:
        add_para(doc, text)
    add_para(doc, "Supporting information captions for second-revision materials", bold_label="").runs[0].bold = True
    for text in [
        "S1 Table. Case-flow and data-availability audit for the second-revision internal-validation analyses.",
        "S2 Table. Internal cross-validated performance of optimized classical clinical baseline models.",
        "S3 Table. Repeated stratified 5-fold cross-validation robustness summary for optimized classical clinical baselines.",
        "S1 Fig. Internal cross-validated calibration curves for optimized classical clinical baseline models.",
        "S2 Fig. Internal decision-curve analysis for optimized classical clinical baseline models.",
        "S3 Fig. Receiver operating characteristic curves for optimized classical clinical baseline models under internal cross-validation.",
        "S4 Fig. Precision-recall curves for optimized classical clinical baseline models under internal cross-validation.",
    ]:
        add_para(doc, text)
    doc.save(str(support_out))


def write_response(response_out: Path, response_md: Path) -> None:
    sections = build_response_sections()
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading("Response to Reviewers - Second Revision", level=0)
    for line in [
        "Manuscript title: Few-Shot Prediction of Post-Fusion Acute Kidney Injury Using Multimodal Radiomics-Clinical Feature-Based Prototype Networks",
        "Journal: PLOS ONE",
        "Manuscript ID: PONE-D-25-38340",
        "Authors: Jing Lan Zhang, Zhen Yang, Jun Li",
    ]:
        add_para(doc, line)
    add_para(doc, "Dear Editor and Reviewer,")
    add_para(
        doc,
        "We sincerely thank you for the careful second-round review and constructive comments. We have revised the manuscript and supporting materials to address concerns regarding external validity, overfitting, missing intraoperative variables, model complexity, baseline comparisons, calibration, decision-curve analysis, and cautious clinical interpretation. All newly added analyses are internal validation analyses and are not presented as external validation or evidence of clinical deployment readiness.",
    )
    for section in sections:
        doc.add_heading(section["title"], level=1)
        add_para(doc, section["comment"], bold_label="Comment: ")
        first = True
        for part in section["response"].split("\n\n"):
            add_para(doc, part, bold_label="Response: " if first else None)
            first = False
        add_para(doc, section["actions"], bold_label="Actions taken: ")
        add_para(doc, section["location"], bold_label="Location in revised manuscript/supporting materials: ")
    add_para(
        doc,
        "We hope that these revisions address the reviewer’s concerns. We have intentionally kept the interpretation conservative and have highlighted that external, temporal, and prospective validation are required before any clinical deployment.",
    )
    add_para(doc, "Sincerely,")
    add_para(doc, "The Authors")
    doc.save(str(response_out))

    md = [
        "# Response to Reviewers - Second Revision",
        "",
        "**Manuscript title:** Few-Shot Prediction of Post-Fusion Acute Kidney Injury Using Multimodal Radiomics-Clinical Feature-Based Prototype Networks",
        "**Journal:** PLOS ONE",
        "**Manuscript ID:** PONE-D-25-38340",
        "",
        "Dear Editor and Reviewer,",
        "",
        "We sincerely thank you for the careful second-round review and constructive comments. We have revised the manuscript and supporting materials to address concerns regarding external validity, overfitting, missing intraoperative variables, model complexity, baseline comparisons, calibration, decision-curve analysis, and cautious clinical interpretation. All newly added analyses are internal validation analyses and are not presented as external validation or evidence of clinical deployment readiness.",
        "",
    ]
    for section in sections:
        md += [
            f"## {section['title']}",
            "",
            f"**Comment:** {section['comment']}",
            "",
            f"**Response:** {section['response']}",
            "",
            f"**Actions taken:** {section['actions']}",
            "",
            f"**Location in revised manuscript/supporting materials:** {section['location']}",
            "",
        ]
    response_md.write_text("\n".join(md), encoding="utf-8")


def write_change_log(change_log: Path) -> None:
    change_log.write_text(
        """# Second Revision Integration Change Log

Generated: 2026-06-14

## Output Files

- `Response_to_Reviewers_Second_Revision.docx`: formal second-revision response letter.
- `Manuscript_Second_Revision_Clean.docx`: clean revised manuscript generated from the clean first-revision manuscript.
- `Supporting_information_Second_Revision.docx`: supporting information with added second-revision baseline/calibration/DCA appendix and captions.
- `Response_to_Reviewers_Second_Revision.md`: editable markdown mirror of the response letter.
- `Strict_Consistency_Audit.md`: automated consistency checks.

## Main Manuscript Changes

- Rewrote the Abstract to define the study as single-center, retrospective, exploratory, and internally validated.
- Clarified that the final multimodal model AUC = 0.816 is an internal cross-validation result.
- Added the newly verified classical clinical baseline comparison: gradient boosting AUC 0.636 (95% CI 0.548-0.718), Brier 0.223, MCC 0.172; repeated CV mean AUC 0.628 +/- 0.084.
- Added calibration and decision-curve analysis text for optimized classical clinical baselines only.
- Added explicit limitation language for missing intraoperative variables: warm ischemia time, pneumoperitoneum pressure/duration, intraoperative hypotension, blood loss, operative duration, fluid balance, vasopressor exposure, and nephrotoxic exposure.
- Tempered SHAP language to descriptive model-attribution wording.
- Added S1-S3 Table and S1-S4 Fig supporting-information captions.

## Interpretation Guardrails

- Do not describe any new analysis as external validation.
- Do not claim final multimodal-model calibration or DCA unless final multimodal probabilities are separately analyzed.
- Do not state that clinical deployment or proven clinical utility has been demonstrated.
- Distinguish the main 194-patient multimodal cohort from the 193-patient tabular classical-baseline analytic set.
""",
        encoding="utf-8",
    )


def write_audit(audit_out: Path, response_out: Path, manuscript_out: Path, support_out: Path) -> None:
    checks = []
    for name, path in [
        ("manuscript", manuscript_out),
        ("response", response_out),
        ("supporting_info", support_out),
    ]:
        text = doc_text(path)
        checks.append(f"## {name}: {path.name}")
        for phrase in [
            "external validation was performed",
            "clinically deployable",
            "proves clinical utility",
            "proof of clinical utility",
            "ready for clinical deployment",
            "successfully avoids overfitting",
        ]:
            found = phrase.lower() in text.lower()
            checks.append(f"- Forbidden phrase check `{phrase}`: {'FOUND' if found else 'not found'}")
        checks.append(
            f"- Mentions internal validation/cross-validation: {'yes' if ('internal' in text.lower() and 'cross-validation' in text.lower()) else 'CHECK MANUALLY'}"
        )
        checks.append("")
    checks.append("## Required files present")
    for path in [
        response_out,
        manuscript_out,
        support_out,
        ROOT / "second_revision_submission_package" / "supplementary_materials_manifest.md",
    ]:
        checks.append(f"- {path}: {'present' if path.exists() else 'MISSING'}")
    audit_out.write_text("\n".join(checks), encoding="utf-8")


def main() -> None:
    OUT.mkdir(exist_ok=True)

    manuscript_src = next(p for p in ROOT.rglob("Manuscript.docx") if p.name == "Manuscript.docx")
    support_src = ROOT / "Supporting information.docx"

    manuscript_out = OUT / "Manuscript_Second_Revision_Clean.docx"
    support_out = OUT / "Supporting_information_Second_Revision.docx"
    response_out = OUT / "Response_to_Reviewers_Second_Revision.docx"
    response_md = OUT / "Response_to_Reviewers_Second_Revision.md"
    change_log = OUT / "Second_Revision_Change_Log.md"
    audit_out = OUT / "Strict_Consistency_Audit.md"

    shutil.copy2(manuscript_src, manuscript_out)
    shutil.copy2(support_src, support_out)

    revise_manuscript(manuscript_out)
    revise_supporting_info(support_out)
    write_response(response_out, response_md)
    write_change_log(change_log)
    write_audit(audit_out, response_out, manuscript_out, support_out)

    print(f"Created outputs in {OUT}")
    for path in [response_out, manuscript_out, support_out, response_md, change_log, audit_out]:
        print(f"{path.name}\t{path.stat().st_size}")


if __name__ == "__main__":
    main()
